import docx
from parsers.base_parser import BaseParser
import os

class DOCXParser(BaseParser):
    def parse(self, file_path: str) -> dict:
        filename = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
                
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    full_text.append(" | ".join(row_data))
                    
        combined_text = "\n".join(full_text)
        cleaned_text = self._clean_text(combined_text)
        
        return {
            'text': cleaned_text,
            'metadata': {
                'filename': filename,
                'pages': 1, # DOCX pagination is dynamic
                'language': 'en',
                'file_size': file_size
            },
            'pages': [{'page_num': 1, 'text': cleaned_text}]
        }
